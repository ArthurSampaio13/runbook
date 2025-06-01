#!/usr/bin/env python3
import boto3
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import requests
from io import BytesIO
from botocore.exceptions import ClientError, NoCredentialsError
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AWSRunbookGenerator:
    def __init__(self, aws_accounts=None):
        self.aws_accounts = aws_accounts or [{'profile_name': 'default', 'region': 'us-east-1'}]
        self.document = Document()
        self.aws_sessions = []
        self.aws_region_image_url = "https://d1.awsstatic.com/legal/aws-logo-web-300x180_563e6ffa42b4e85e2a6fd8bda7e7a7b6c84efbe7.png"
        
    def setup_aws_sessions(self):
        for account in self.aws_accounts:
            try:
                if 'role_arn' in account:
                    sts_client = boto3.client('sts')
                    assumed_role = sts_client.assume_role(
                        RoleArn=account['role_arn'],
                        RoleSessionName='RunbookGenerator'
                    )
                    
                    session = boto3.Session(
                        aws_access_key_id=assumed_role['Credentials']['AccessKeyId'],
                        aws_secret_access_key=assumed_role['Credentials']['SecretAccessKey'],
                        aws_session_token=assumed_role['Credentials']['SessionToken'],
                        region_name=account.get('region', 'us-east-1')
                    )
                else:
                    session = boto3.Session(
                        profile_name=account.get('profile_name', 'default'),
                        region_name=account.get('region', 'us-east-1')
                    )
                
                self.aws_sessions.append({
                    'session': session,
                    'account_config': account
                })
                
            except Exception as e:
                logger.error(f"Failed to setup session for account {account}: {str(e)}")
                
    def create_cover_page(self):
        title = self.document.add_heading('AWS Runbook', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.document.add_paragraph()
        subtitle_run = subtitle.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
        subtitle_run.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_page_break()
        
    def create_table_of_contents(self):
        toc_heading = self.document.add_heading('Table of Contents', level=1)
        
        toc_data = [
            ('Summary', '2'),
            ('Global Positioning', '3'),
            ('Account Summary', '4'),
            ('AWS Organizations', '4'),
            ('AWS Control Tower', '4'),
            ('VPC Summary', '5'),
            ('Public DNS (Route 53)', '9'),
            ('EC2 Instances', '9'),
            ('RDS', '9'),
            ('API Gateway', '10'),
            ('CloudFront', '10'),
            ('Lambda', '10'),
            ('SNS', '11'),
            ('AWS Backup', '11'),
            ('EventBridge', '12'),
            ('LoadBalancer', '12'),
            ('ECS', '13'),
            ('IAM Identity Center (SSO)', '13')
        ]
        
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Section'
        hdr_cells[1].text = 'Page'
        
        for section, page in toc_data:
            row_cells = table.add_row().cells
            row_cells[0].text = section
            row_cells[1].text = page
            
        self.document.add_page_break()
        
    def add_aws_region_image(self):
        try:
            response = requests.get(self.aws_region_image_url)
            if response.status_code == 200:
                image_stream = BytesIO(response.content)
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(4))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                self.document.add_paragraph("AWS Region Image (Unable to load)")
        except Exception as e:
            logger.error(f"Failed to add AWS region image: {str(e)}")
            self.document.add_paragraph("AWS Region Image (Unable to load)")
            
    def get_account_summary(self):
        account_data = []
        
        for aws_session in self.aws_sessions:
            session = aws_session['session']
            try:
                sts_client = session.client('sts')
                account_id = sts_client.get_caller_identity()['Account']
                
                try:
                    support_client = session.client('support', region_name='us-east-1')
                    support_plan = "Basic" 
                except:
                    support_plan = "Unknown"
                
                try:
                    iam_client = session.client('iam')
                    aliases = iam_client.list_account_aliases()
                    account_alias = aliases['AccountAliases'][0] if aliases['AccountAliases'] else "No alias"
                except:
                    account_alias = "Unknown"
                
                ec2_client = session.client('ec2')
                regions = ec2_client.describe_regions()
                active_regions = len(regions['Regions'])
                
                account_data.append({
                    'account_id': account_id,
                    'support_plan': support_plan,
                    'account_alias': account_alias,
                    'active_regions': str(active_regions),
                    'account_role': aws_session['account_config'].get('role_arn', 'Direct Access')
                })
                
            except Exception as e:
                logger.error(f"Failed to get account summary: {str(e)}")
                account_data.append({
                    'account_id': 'Unknown',
                    'support_plan': 'Unknown',
                    'account_alias': 'Unknown',
                    'active_regions': 'Unknown',
                    'account_role': 'Unknown'
                })
                
        return account_data
        
    def get_vpc_summary(self):
        vpc_data = []
        
        for aws_session in self.aws_sessions:
            session = aws_session['session']
            try:
                ec2_client = session.client('ec2')
                vpcs = ec2_client.describe_vpcs()
                
                for vpc in vpcs['Vpcs']:
                    vpc_name = 'N/A'
                    if 'Tags' in vpc:
                        for tag in vpc['Tags']:
                            if tag['Key'] == 'Name':
                                vpc_name = tag['Value']
                                break
                    
                    vpc_data.append({
                        'vpc_id': vpc['VpcId'],
                        'vpc_name': vpc_name,
                        'cidr_block': vpc['CidrBlock'],
                        'state': vpc['State'],
                        'is_default': str(vpc['IsDefault'])
                    })
                    
            except Exception as e:
                logger.error(f"Failed to get VPC summary: {str(e)}")
                
        return vpc_data
        
    def get_ec2_instances(self):
        ec2_data = []
        
        for aws_session in self.aws_sessions:
            session = aws_session['session']
            try:
                ec2_client = session.client('ec2')
                instances = ec2_client.describe_instances()
                
                for reservation in instances['Reservations']:
                    for instance in reservation['Instances']:
                        instance_name = 'N/A'
                        if 'Tags' in instance:
                            for tag in instance['Tags']:
                                if tag['Key'] == 'Name':
                                    instance_name = tag['Value']
                                    break
                        
                        ec2_data.append({
                            'instance_id': instance['InstanceId'],
                            'instance_name': instance_name,
                            'instance_type': instance['InstanceType'],
                            'state': instance['State']['Name'],
                            'private_ip': instance.get('PrivateIpAddress', 'N/A'),
                            'public_ip': instance.get('PublicIpAddress', 'N/A')
                        })
                        
            except Exception as e:
                logger.error(f"Failed to get EC2 instances: {str(e)}")
                
        return ec2_data
        
    def get_rds_instances(self):
        rds_data = []
        
        for aws_session in self.aws_sessions:
            session = aws_session['session']
            try:
                rds_client = session.client('rds')
                db_instances = rds_client.describe_db_instances()
                
                for db in db_instances['DBInstances']:
                    rds_data.append({
                        'db_instance_id': db['DBInstanceIdentifier'],
                        'db_engine': db['Engine'],
                        'db_class': db['DBInstanceClass'],
                        'status': db['DBInstanceStatus'],
                        'endpoint': db.get('Endpoint', {}).get('Address', 'N/A'),
                        'port': str(db.get('Endpoint', {}).get('Port', 'N/A'))
                    })
                    
            except Exception as e:
                logger.error(f"Failed to get RDS instances: {str(e)}")
                
        return rds_data
        
    def get_lambda_functions(self):
        lambda_data = []
        
        for aws_session in self.aws_sessions:
            session = aws_session['session']
            try:
                lambda_client = session.client('lambda')
                functions = lambda_client.list_functions()
                
                for func in functions['Functions']:
                    lambda_data.append({
                        'function_name': func['FunctionName'],
                        'runtime': func['Runtime'],
                        'handler': func['Handler'],
                        'code_size': str(func['CodeSize']),
                        'timeout': str(func['Timeout']),
                        'memory_size': str(func['MemorySize'])
                    })
                    
            except Exception as e:
                logger.error(f"Failed to get Lambda functions: {str(e)}")
                
        return lambda_data
        
    def create_data_table(self, title, headers, data):
        self.document.add_heading(title, level=2)
        
        if not data:
            self.document.add_paragraph("No data available")
            return
            
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for item in data:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                key = header.lower().replace(' ', '_')
                row_cells[i].text = str(item.get(key, 'N/A'))
                
    def generate_runbook(self):
        logger.info("Starting AWS Runbook generation...")
        
        self.setup_aws_sessions()
        
        self.create_cover_page()
        
        self.create_table_of_contents()
        
        self.document.add_heading('Summary', level=1)
        self.document.add_paragraph(
            "This runbook provides a comprehensive overview of AWS resources across "
            "the specified accounts. It includes information about compute, storage, "
            "networking, and security services."
        )
        
        self.document.add_heading('Global Positioning', level=1)
        self.document.add_paragraph(
            "This section provides an overview of the AWS global infrastructure "
            "and regional positioning for the analyzed accounts."
        )
        self.add_aws_region_image()
        
        account_data = self.get_account_summary()
        self.create_data_table(
            'Account Summary',
            ['Account ID', 'Support Plan', 'Account Alias', 'Active Regions', 'Account Role'],
            account_data
        )
        
        self.create_data_table('AWS Organizations', ['Organization ID', 'Master Account', 'Status'], [])
        
        self.create_data_table('AWS Control Tower', ['Landing Zone', 'Status', 'Version'], [])
        
        vpc_data = self.get_vpc_summary()
        self.create_data_table(
            'VPC Summary',
            ['VPC ID', 'VPC Name', 'CIDR Block', 'State', 'Is Default'],
            vpc_data
        )
        
        self.create_data_table('Public DNS (Route 53)', ['Hosted Zone', 'Domain Name', 'Type', 'Records'], [])
        
        ec2_data = self.get_ec2_instances()
        self.create_data_table(
            'EC2 Instances',
            ['Instance ID', 'Instance Name', 'Instance Type', 'State', 'Private IP', 'Public IP'],
            ec2_data
        )
        
        rds_data = self.get_rds_instances()
        self.create_data_table(
            'RDS',
            ['DB Instance ID', 'DB Engine', 'DB Class', 'Status', 'Endpoint', 'Port'],
            rds_data
        )
        
        self.create_data_table('API Gateway', ['API ID', 'API Name', 'Type', 'Stage', 'Endpoint'], [])
        
        self.create_data_table('CloudFront', ['Distribution ID', 'Domain Name', 'Status', 'Origin'], [])
        
        lambda_data = self.get_lambda_functions()
        self.create_data_table(
            'Lambda',
            ['Function Name', 'Runtime', 'Handler', 'Code Size', 'Timeout', 'Memory Size'],
            lambda_data
        )
        
        self.create_data_table('SNS', ['Topic ARN', 'Topic Name', 'Subscriptions', 'Protocol'], [])
        
        self.create_data_table('AWS Backup', ['Backup Plan', 'Backup Vault', 'Recovery Points', 'Status'], [])
        
        self.create_data_table('EventBridge', ['Rule Name', 'Event Pattern', 'Targets', 'Status'], [])
        
        self.create_data_table('LoadBalancer', ['Load Balancer Name', 'Type', 'Scheme', 'State', 'DNS Name'], [])
        
        self.create_data_table('ECS', ['Cluster Name', 'Service Name', 'Task Definition', 'Running Tasks', 'Status'], [])
        
        self.create_data_table('IAM Identity Center (SSO)', ['Instance ARN', 'Identity Source', 'Status', 'Users'], [])
        
        filename = f"aws_runbook_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        self.document.save(filename)
        logger.info(f"Runbook generated successfully: {filename}")
        
        return filename

def main():
    aws_accounts = [
         # {
-        #     'account_id': '123456789012',
-        #     'role_arn': 'arn:aws:iam::123456789012:role/CrossAccountRole',
-        #     'region': 'us-west-2'
-        # } 
        {
            'profile_name': 'default',
            'region': 'us-east-1'
        },

        ]
    
    generator = AWSRunbookGenerator(aws_accounts)
    filename = generator.generate_runbook()
    print(f"AWS Runbook generated: {filename}")

if __name__ == "__main__":
    main()